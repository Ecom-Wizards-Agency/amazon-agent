"""Branded client-plan DOCX renderer for the Amazon launch strategy model."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "amazon-ad-audit"))
import branding as _branding  # noqa: E402

# These values were correct by luck, not by construction: they were hand-copied and happened
# to match. Load them instead, so a brand change reaches this builder like every other.
# Contract: company-ai-skills/skills/ecom-wizards-brand/.
_BRAND = _branding.load_branding({})
_DOC = _BRAND["palette_doc"]
INK = _DOC["ink"]
ACCENT = _DOC["accent"]
CLOUD = _DOC["cloud"]
MISTLINE = _DOC["mistline"]
STEEL = _DOC["steel"]
MIST = _DOC["mist"]
WHITE = _DOC["white"]
FONT = _BRAND["fonts"]["doc_font_name"]
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, color: str = INK, bold: bool | None = None, italic: bool | None = None):
    run.font.name = FONT
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MISTLINE, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8, color=MIST)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=INK):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(value)
        set_run_font(run, size=8.2, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, CLOUD)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(value)
            set_run_font(run, size=8.2, color=INK)
    return table


def add_callout(doc, label: str, text: str, color=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color=MISTLINE, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CLOUD)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label.upper())
    set_run_font(run, size=8, color=color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run(text)
    set_run_font(run, size=10.2, color=INK)
    return table


def add_body(doc, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=9.7, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=9.7)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.7)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(" " + text)
    set_run_font(run, size=9.5)
    return p


def fmt_money(value: Any, currency: str) -> str:
    return "N/A" if value is None else f"{currency} {float(value):,.0f}"


def fmt_pct(value: Any) -> str:
    return "N/A" if value is None else f"{float(value) * 100:.1f}%"


def fmt_units(value: Any) -> str:
    return "Unconfirmed" if value is None else f"{float(value):,.0f}"


def _configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in (
        ("Heading 1", 16, 16, 7),
        ("Heading 2", 12.5, 12, 5),
        ("Heading 3", 10.5, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(INK if name == "Heading 1" else STEEL)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_header_footer(section, client_name: str, logo_path: Path):
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    def populate_header(header):
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        set_table_geometry(table, [3000, 6360], indent=0)
        set_table_borders(table, color=WHITE, size=0)
        left = table.cell(0, 0)
        left.paragraphs[0].add_run().add_picture(str(logo_path), width=Inches(1.25))
        right = table.cell(0, 1)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = right.paragraphs[0].add_run(f"90-DAY LAUNCH PLAN | {client_name}".upper())
        set_run_font(run, size=7.4, color=MIST, bold=True)

    def populate_footer(footer):
        ftable = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        set_table_geometry(ftable, [3400, 2560, 3400], indent=0)
        set_table_borders(ftable, color=WHITE, size=0)
        values = ["Amazon Launch Strategy", None, "www.ecomwizards.agency"]
        for idx, value in enumerate(values):
            p = ftable.cell(0, idx).paragraphs[0]
            p.alignment = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT][idx]
            if value:
                run = p.add_run(value)
                set_run_font(run, size=7.2, color=MIST)
            else:
                run = p.add_run("Page ")
                set_run_font(run, size=7.2, color=MIST)
                add_field(p, "PAGE")
                run = p.add_run(" of ")
                set_run_font(run, size=7.2, color=MIST)
                add_field(p, "NUMPAGES")

    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def build_document(model: dict[str, Any], output_path: str | Path) -> Path:
    config = model["generated_from"]
    client = config["client"]
    currency = client["currency"]
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    repo_root = Path(__file__).resolve().parents[2]
    logo_path = repo_root / "tools/amazon-ad-audit/brand/logo_black.png"

    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    _configure_styles(doc)
    _add_header_footer(section, client["brand"], logo_path)

    commercial = model.get("commercial")
    if not commercial:
        raise ValueError("The executive document requires commercial_targets in the launch config.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("EXECUTIVE AMAZON LAUNCH BRIEF")
    set_run_font(run, size=8, color=ACCENT, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(f"{client['brand']} | 90-Day Commercial Plan")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(11)
    run = subtitle.add_run(f"{client['account']} | {client['marketplace']} | {client.get('launch_timing_label', 'Day 0 and Weeks 1-13')}")
    set_run_font(run, size=10.5, color=STEEL)

    add_callout(
        doc,
        "Executive objective",
        "Build from launch to a $300/day Month 1 exit, reach $1,000/day by the end of Month 2, maintain $1,000/day as the Month 3 commitment, and scale toward $2,000/day only when conversion, margin, and stock gates are met.",
    )
    doc.add_heading("1. Objective and 90-day milestones", level=1)
    milestone_rows = [
        ["Month 1", "$0 to $300/day", "$4,200", "Prove relevance and conversion"],
        ["Month 2", "$300 to $1,000/day", "$18,200", "Scale proven non-branded demand"],
        ["Month 3 committed", "Hold $1,000/day", "$35,000", "Stabilize the repeatable base"],
        ["Month 3 stretch", "$1,000 to $2,000/day", "$52,500", "Unlock only after scaling gates"],
    ]
    add_table(doc, ["Period", "Revenue run-rate", "Period sales", "Operating objective"], milestone_rows, [1800, 1900, 1550, 4110])
    add_body(doc, "These are operating objectives, not guaranteed sales. The launch date remains Day 0 until confirmed.")

    doc.add_heading("Commercial outcome", level=1)
    summary_rows = [[
        item["path"],
        fmt_money(item["target_revenue"], currency),
        fmt_units(item["forecast_units"]),
        fmt_units(item["customer_sale_inventory_required"]),
        item["month_3_objective"],
    ] for item in commercial["summaries"]]
    add_table(doc, ["Path", "13-week sales", "Forecast units", "Stock incl. 20%", "Month 3"], summary_rows, [1750, 1750, 1450, 1650, 2760])
    add_callout(doc, "Recommended commitment", "Fund and stock the committed path first. Keep stretch funding available, but release it only when conversion, economics, and inventory coverage prove the next step.")

    doc.add_page_break()
    doc.add_heading("2. Investment, expected sales, and scaling gates", level=1)
    ppc_order = ("month_1", "month_2", "month_3_committed", "month_3_stretch")
    ppc_rows = []
    for plan_id in ppc_order:
        plan = commercial["ppc_plan"][plan_id]
        ppc_rows.append([plan["label"], fmt_money(plan["planned_spend"], currency), fmt_money(plan["spend_ceiling"], currency), "Available, not forced"])
    add_table(doc, ["Period", "Planned PPC", "Funding ceiling", "Control"], ppc_rows, [2100, 1750, 1800, 3710])
    add_body(doc, "Planned spend is the working budget. The ceiling is approved capacity, not a requirement to spend. Actual pacing is governed by search-term relevance, conversion, margin visibility, and inventory cover.")
    add_bullet(doc, "Month 1 gate: search terms are relevant, listings convert, and stock can support the $300/day exit rate.")
    add_bullet(doc, "Month 2 gate: proven targets can absorb more budget without losing conversion quality or creating a stockout risk.")
    add_bullet(doc, "Month 3 stretch gate: landed economics are confirmed, the committed $1,000/day level is stable, and inventory covers the upside path.")

    doc.add_heading("3. Where the PPC money goes", level=1)
    allocation_rows = []
    for plan_id in ppc_order:
        plan = commercial["ppc_plan"][plan_id]
        alloc = plan["campaign_allocation"]
        allocation_rows.append([
            plan["label"],
            fmt_pct(alloc["high_intent_non_branded"]),
            fmt_pct(alloc["discovery"]),
            fmt_pct(alloc["competitor_keywords"]),
            fmt_pct(alloc["competitor_product_targeting"]),
            fmt_pct(alloc["branded_defense"]),
        ])
    add_table(doc, ["Period", "Core non-brand", "Discovery", "Comp. KW", "Comp. PT", "Brand"], allocation_rows, [1850, 1800, 1300, 1400, 1500, 1510])
    keywords = commercial.get("keywords", {})
    add_bullet(doc, "Core non-branded: " + ", ".join(keywords.get("core", [])) + ".")
    add_bullet(doc, "Discovery: " + ", ".join(keywords.get("discovery", [])) + ".")
    add_bullet(doc, "Competitor direction: " + ", ".join(keywords.get("competitors", [])) + ".")
    add_bullet(doc, "Controlled test only: " + ", ".join(keywords.get("controlled_tests", [])) + ". " + keywords.get("controlled_test_guardrail", ""))

    doc.add_page_break()
    doc.add_heading("4. Stock requirement and configuration split", level=1)
    products = {product["id"]: product for product in config["products"] if product.get("phase") == "launch"}
    stock_rows = []
    for item in commercial["summaries"]:
        for product_id, units in item["customer_stock_by_product"].items():
            vine_units = int(products[product_id].get("reviews", {}).get("vine_units") or 0)
            stock_rows.append([
                item["path"],
                products[product_id]["name"],
                fmt_units(item["product_units"][product_id]),
                fmt_units(units),
                fmt_units(vine_units),
                fmt_units(item["total_stock_by_product"][product_id]),
            ])
    add_table(doc, ["Path", "Configuration", "Forecast sales", "+20% stock", "Vine", "Total required"], stock_rows, [1450, 2200, 1450, 1450, 1000, 1810])
    add_body(doc, "The phased mix is 85/15 Starter/Refill in Month 1, 75/25 in Month 2, and 65/35 in Month 3. This produces an approximately 70/30 blended launch mix.")
    add_callout(doc, "Inventory decision", "Reserve 675 units for the committed path or 875 units for the stretch path, plus any confirmed Vine allocation. The 1,075-unit capacity ceiling applies only if $2,000/day is held throughout Month 3.")
    add_body(doc, "No dependable replenishment is assumed during Weeks 1-13 until opening stock, inbound units, MOQ, production time, freight time, and FBA receiving buffer are confirmed.")

    doc.add_heading("Offer scope", level=1)
    launch_rows = []
    for product in config["products"]:
        price = product.get("launch_price") if product.get("launch_price") is not None else product.get("list_price")
        launch_rows.append([product["name"], "Launch" if product.get("phase") == "launch" else "Later phase", fmt_money(price, currency), product.get("phase_gate", "Day 0")])
    add_table(doc, ["Offer", "Phase", "Price", "Gate"], launch_rows, [2400, 1450, 1350, 4160])

    doc.add_page_break()
    doc.add_heading("Appendix | Decisions, controls, and open confirmations", level=1)
    doc.add_heading("Pricing and reviews", level=2)
    add_bullet(doc, "Starter Kit price: $104.99. Refill Pouch price: $99.99. Hold these prices until landed COGS, Amazon fees, contribution margin, and the discount floor are confirmed.")
    add_bullet(doc, "Use Vine only after eligibility is confirmed and add the approved allocation above customer-sale inventory.")
    add_bullet(doc, "Use Amazon's standard Request a Review flow. Helium 10 Follow-Up may automate only Amazon's standard template with one request per order and deterministic deduplication.")
    add_bullet(doc, "No incentives, review gating, creator purchases for reviews, disguised compensation, or duplicate requests.")
    add_bullet(doc, "External halo remains zero until measured evidence or an explicit editable assumption is supplied.")

    doc.add_heading("Owners and open confirmations", level=2)
    owner_rows = [[item.get("decision", ""), item.get("owner", "Unassigned"), item.get("status", "Open")] for item in config.get("owners", [])]
    add_table(doc, ["Decision", "Owner", "Status"], owner_rows, [4800, 2800, 1760])
    grouped_missing: dict[str, list[str]] = {}
    for item in model["validation"]["missing"]:
        label = item["label"]
        if ": " in label:
            group, detail = label.split(": ", 1)
        elif "Meta" in label or "Google" in label or "Branded-search" in label:
            group, detail = "External support", label
        elif "revenue" in label.lower() or "orders" in label.lower():
            group, detail = "Current Amazon baseline", label
        else:
            group, detail = "Launch governance", label
        grouped_missing.setdefault(group, []).append(detail)
    for group, details in grouped_missing.items():
        add_bullet(doc, f"OPEN, {group}: {'; '.join(details)}")
    add_callout(doc, "Day 0 gate", "Do not move beyond planning until listing and fulfillment readiness, stock coverage, margin floor, PPC funding, review eligibility, external support, and launch timing are confirmed by the named owners.")

    doc.core_properties.title = f"{client['brand']} 90-Day Amazon Launch Strategy"
    doc.core_properties.subject = "PPC, pricing, inventory, reviews, and launch forecast"
    doc.core_properties.author = "Ecom Wizards"
    doc.save(output_path)
    return output_path


def effective_price(product: dict[str, Any]) -> float:
    launch = product.get("launch_price")
    if launch is not None:
        return float(launch)
    price = float(product.get("list_price") or 0)
    return price * (1 - float(product.get("coupon_pct") or 0))
