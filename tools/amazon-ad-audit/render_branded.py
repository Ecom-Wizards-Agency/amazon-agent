#!/usr/bin/env python3
"""Branded audit renderer: narrative markdown + metrics -> A4 .docx.

Client- and agency-agnostic. Consumes the operator-written `*_Sales_Audit_SCAFFOLD.md` (the narrative),
the run's `metrics.json` (for the KPI cards), the config (client / market / window / prepared_by /
first_time), the agency identity from `_local/branding/branding.json` (see BRANDING.md; falls back to a
agent-neutral example style), and the local brand assets in `brand/`. Produces a light, readable body,
a dark **cover page** (first-time audits only), KPI stat-cards, restyled tables, a full-lockup running
header, a text-only three-part footer, and smart page-break hygiene. Degrades gracefully to plain
md_to_docx when assets are missing.

Markdown conventions (superset of md_to_docx.py — which has no image support):
  ## H2                      -> orange rule + section head
  ### Lever N: title  /  **Lever N — title**  -> orange "LEVER N" eyebrow + title
  ### H3                     -> sub-head
  > quote                    -> orange note callout
  ![caption](file.png)       -> figure + caption (path relative to the markdown file)
  | a | b |                  -> table (Ink header, tabular numbers)
  - / *  |  1.               -> bullets / numbered
  <!-- ... -->               -> dropped
KPI cards are auto-built from metrics.json and inserted right after the first H2.
"""
from __future__ import annotations
import re, json
from pathlib import Path

HERE = Path(__file__).resolve().parent

import branding as _branding
from analyze_audit import blended_metrics_available

# palette/fonts — populated from the branding file (agency identity) via _apply_branding()
INK_H = CLOUD_H = ORANGE_H = MISTLINE_H = STEEL_H = MIST_H = ""
FONT_NAME = FONT_FILE = ""
_B: dict = {}


def _apply_branding(b):
    global _B, INK_H, CLOUD_H, ORANGE_H, MISTLINE_H, STEEL_H, MIST_H, FONT_NAME, FONT_FILE
    _B = b
    p = b["palette_doc"]
    INK_H, CLOUD_H, ORANGE_H = p["ink"], p["cloud"], p["accent"]
    MISTLINE_H, STEEL_H, MIST_H = p["mistline"], p["steel"], p["mist"]
    FONT_NAME, FONT_FILE = b["fonts"]["doc_font_name"], b["fonts"]["doc_font_file"]


_apply_branding(_branding.load_branding({}))


def _money(v, cur="USD"):
    s = "€" if cur in ("EUR", "€") else "$"
    if abs(v) >= 1_000_000: return f"{s}{v/1_000_000:.1f}M"
    if abs(v) >= 1_000: return f"{s}{round(v/1000)}K"
    return f"{s}{v:,.0f}"


def _bcfg(cfg, key, default=None):
    """Read a branding field from cfg['branding'][key], falling back to top-level, then default."""
    b = cfg.get("branding", {}) or {}
    v = b.get(key, cfg.get(key))
    return default if v in (None, "") else v


def _doc_label(cfg):
    return _bcfg(cfg, "doc_label", "Account Audit")


def _footer_page_words(cfg):
    """Footer page counter wording. Defaults to English so existing renders are unchanged;
    set branding.footer_page_words to ["Seite ", " von "] for a German deliverable."""
    w = _bcfg(cfg, "footer_page_words", None)
    if isinstance(w, (list, tuple)) and len(w) == 2:
        return str(w[0]), str(w[1])
    return "page ", " of "


def _report_month(cfg, M):
    """Return the deliverable month/year for the running header."""
    candidates = [str(cfg.get("date", ""))]
    windows = M.get("windows", {}) or cfg.get("windows", {}) or {}
    candidates.extend(str(windows.get(k, "")) for k in ("ads", "business_report"))
    for raw in candidates:
        dates = re.findall(r'\b(20\d{2})-(\d{2})-(\d{2})\b', raw)
        if dates:
            year, month, _ = dates[-1]
            names = ("", "JANUARY", "FEBRUARY", "MARCH", "APRIL", "MAY", "JUNE",
                     "JULY", "AUGUST", "SEPTEMBER", "OCTOBER", "NOVEMBER", "DECEMBER")
            mi = int(month)
            if 1 <= mi <= 12:
                return f"{names[mi]} {year}"
    return ""


def _running_text(cfg, M):
    label = _doc_label(cfg)
    month = _report_month(cfg, M)
    return {
        "header_right": f"{label.upper()} · {month}" if month else label.upper(),
        "footer_left": f"{label} · {cfg.get('client', '')}".rstrip(" ·"),
        "footer_right": _B.get("agency_url", ""),
    }


def _kpi_after(blocks):
    """Block index of the H2 to insert the KPI cards after: the verdict/summary section, else the first H2."""
    h2s = [(i, p) for i, (k, p) in enumerate(blocks) if k == "h2"]
    for i, p in h2s:
        if re.search(r'verdict|summary|snapshot|performance|the numbers', p, re.I): return i
    return h2s[0][0] if h2s else -1


def _kpis(M):
    if M.get("custom_kpis"):  # non-audit docs: [[number, label, sub-or-null], ...] in metrics.json
        return [(str(n), l, s) for n, l, s in M["custom_kpis"]]
    T = M["totals"]; cur = M.get("currency", "USD"); be = M.get("breakeven", 0.5)
    # Blended cards exist only when Ads and Business Report cover the same dates. When
    # they do not, use ROAS from the Ads snapshot rather than putting a false TACoS beside
    # two incompatible periods.
    fourth = ((f"{T['tacos']*100:.0f}%", "TACoS", None)
              if blended_metrics_available(M)
              else (f"{T['roas']:.2f}x", "Ad ROAS", "same Ads window"))
    spend_label = "Ad spend / snapshot" if M.get("ads_snapshot_directional") else "Ad spend / window"
    return [(_money(T["spend"], cur), spend_label, None),
            (_money(T["sales"], cur), "Ad sales", None),
            (f"{T['acos']*100:.0f}%", "Blended ACoS", f"vs ~{be*100:.0f}% break-even"),
            fourth]


# ------------------------------ markdown -> blocks ------------------------------
# Bold lead-in form: **Lever N — Title.** body text...  -> title = group2, body = group3
_LEVER_INLINE = re.compile(r'^\*\*\s*lever\s+(\d+)\s*[:\-—–]\s*(.+?)\.?\s*\*\*\s*(.*)$', re.I)
# Heading form: ### Lever N: Title   (or a bare "Lever N — Title" line)
_LEVER = re.compile(r'^(?:###\s*)?lever\s+(\d+)\s*[:\-—–]\s*(.+?)\.?$', re.I)
_IMG = re.compile(r'^!\[(.*?)\]\((.*?)\)$')
# Inline markdown the branded renderer understands: **bold** and `code`. Keyword and
# search-term names are written as `code` in the narratives, so without this the
# backticks printed literally in the docx.
_INLINE = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\((https?://[^)]+)\)')
_CODE = re.compile(r'`([^`]+)`')
_COMMENT = re.compile(r'<!--.*?-->', re.S)


# A block construct starts a new block; anything else is a continuation of the line above.
_BLOCK_START = re.compile(r'^\s*(?:#{1,6}\s|\||>\s|[-*]\s|\d+\.\s|!\[|---\s*$|—\s*$|\\pagebreak\s*$)', re.I)


def unwrap_soft_breaks(md):
    """Join hard-wrapped lines back into one paragraph per paragraph.

    The block parser below is line-based, so a prose paragraph wrapped at 100 columns
    used to become one docx paragraph per LINE, each with its own space-after. Worse,
    inline spans are matched per line, so a `**bold**` that straddled a wrap shipped to
    the client as literal asterisks. Four of them reached a delivered audit before this
    was caught, which is exactly the failure mode that leaves no error behind.

    Markdown already says consecutive non-blank text lines are one paragraph, so this
    only makes the renderer agree with the format it claims to read. A blank line, a
    heading, table row, bullet, quote, image or rule still starts a new block.
    """
    out = []
    for line in md.splitlines():
        stripped = line.strip()
        if (out and stripped and out[-1].strip()
                and not _BLOCK_START.match(line)
                and not _BLOCK_START.match(out[-1])):
            out[-1] = out[-1].rstrip() + " " + stripped
        elif (out and stripped and out[-1].strip() and line.startswith(("  ", "\t"))
              and _BLOCK_START.match(out[-1]) and not _BLOCK_START.match(line)):
            # An indented continuation belongs to the bullet or list item above it.
            out[-1] = out[-1].rstrip() + " " + stripped
        else:
            out.append(line)
    return "\n".join(out)


def parse_markdown(md, base_dir):
    md = unwrap_soft_breaks(md)
    blocks = []; tbl = []; title = None
    def flush():
        nonlocal tbl
        if tbl:
            rows = [r for r in tbl if not all(re.match(r'^:?-+:?$', c.strip() or '-') for c in r)]
            if rows: blocks.append(("table", rows))
            tbl = []
    for raw in _COMMENT.sub("", md).splitlines():
        s = raw.rstrip()
        if re.match(r'^\|.*\|\s*$', s):
            tbl.append([c.strip() for c in s.strip().strip('|').split('|')]); continue
        flush()
        if not s.strip(): continue
        if s.strip().lower() == "\\pagebreak": blocks.append(("pagebreak", None)); continue
        # Emit the H1 and the byline as blocks as well as capturing the title. The
        # cover carries both when there is one, and _render_docx drops these blocks in
        # that case. With cover=False nothing else carries them, and the document used
        # to open straight on "Sources:" with no title and no byline anywhere.
        if s.startswith("# "):
            title = s[2:].strip(); blocks.append(("doctitle", title)); continue
        if s.strip() in ("---", "—"): continue
        mLi = _LEVER_INLINE.match(s.strip())
        if mLi:
            blocks.append(("lever", (int(mLi.group(1)), mLi.group(2).strip().rstrip("."))))
            body = mLi.group(3).strip()
            if body: blocks.append(("p", body))
            continue
        mL = _LEVER.match(s.strip())
        if mL: blocks.append(("lever", (int(mL.group(1)), mL.group(2).strip().rstrip(".")))); continue
        if s.startswith("#### "): blocks.append(("h4", s[5:].strip())); continue
        if s.startswith("### "): blocks.append(("h3", s[4:].strip())); continue
        if s.startswith("## "): blocks.append(("h2", s[3:].strip())); continue
        if s.startswith("> "): blocks.append(("note", s[2:].strip())); continue
        mI = _IMG.match(s.strip())
        if mI:
            p = Path(mI.group(2))
            if not p.is_absolute(): p = Path(base_dir) / p
            blocks.append(("img", (str(p), mI.group(1)))); continue
        if re.match(r'^\s*[-*] ', s): blocks.append(("bul", re.sub(r'^\s*[-*] ', '', s))); continue
        if re.match(r'^\s*\d+\. ', s): blocks.append(("num", re.sub(r'^\s*\d+\. ', '', s))); continue
        # a leading "**Prepared by ...**" byline line is cover material when a cover is
        # drawn; _render_docx drops the block in that case and renders it otherwise.
        if re.match(r'^\*\*Prepared by', s.strip()):
            blocks.append(("byline", s.strip().strip("*"))); continue
        blocks.append(("p", s.strip()))
    flush()
    return title, blocks


def _cover_kwargs(cfg, M, blocks, brand_dir):
    client = cfg.get("client", "Client")
    markets = cfg.get("marketplaces", []) or []
    mname = "UNITED STATES" if markets == ["US"] else " / ".join(markets)
    win = (M.get("windows", {}) or cfg.get("windows", {})).get("business_report", "")
    sub = _bcfg(cfg, "cover_subtitle", "A straight look at the account, with recommendations on where to start.")
    words = sub.split(); lines = []; cur = ""
    for w in words:
        if len(cur) + len(w) + 1 > 30: lines.append(cur); cur = w
        else: cur = (cur + " " + w).strip()
    if cur: lines.append(cur)
    # How many H2s the cover's "What's inside" list shows. Default 4 keeps every
    # previously rendered cover byte-identical; raise it via branding.cover_inside_max
    # when the document has more sections worth advertising than the cover currently
    # admits (the block sits in generous whitespace, so 7-8 still fits an A4 cover).
    inside_max = int(_bcfg(cfg, "cover_inside_max", 4) or 4)
    inside = [(f"{i+1:02d}", t) for i, (k, t) in enumerate([b for b in blocks if b[0] == "h2"])
              if t.lower() not in ("sources used", "method notes")][:inside_max]
    by = _bcfg(cfg, "prepared_by", _B.get("prepared_by_default") or "the operator")
    label = _doc_label(cfg)
    return dict(brand_dir=str(brand_dir), title=client,
                eyebrow=label.upper(), dateline=f"{mname} · {win}".upper(),
                sub_lines=lines or [sub], prepared_for=client,
                prepared_by=_branding.prepared_by_line(_B, by),
                inside=inside, footer_left=_branding.cover_footer_left(_B),
                footer_right=label, palette=_B["palette_doc"], font_file=FONT_FILE,
                logo_file=_B["assets"]["logo_white"])


# ------------------------------ DOCX ------------------------------
def _render_docx(blocks, M, cfg, brand_dir, cover_png, out):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    def _rgb(h): return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    INK = _rgb(INK_H); STEEL = _rgb(STEEL_H)
    MIST = _rgb(MIST_H); ORANGE = _rgb(ORANGE_H); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    A4_W = Inches(8.27); A4_H = Inches(11.69)
    doc = Document()
    nm = doc.styles['Normal']; nm.font.name = FONT_NAME; nm.font.size = Pt(10.5); nm.font.color.rgb = INK
    pf = nm.paragraph_format; pf.line_spacing = 1.4; pf.space_after = Pt(7); pf.widow_control = True

    # Google Docs may discard a run-level override when it imports or later promotes a
    # paragraph to a named heading. Brand the named styles themselves so an inherited
    # Heading 2 can never fall back to Google's default blue.
    for style_name, size, color in (
        ('Heading 1', 18, INK),
        ('Heading 2', 12.5, INK),
        ('Heading 3', 11, ORANGE),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    def font(r):
        r.font.name = FONT_NAME; rpr = r._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
        if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT_NAME)

    def runs(p, text, size=10.5, color=INK, bold=False, italic=False, caps=False, tracking=None):
        def emit_link(label, url):
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            rid = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
            h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), rid)
            r = OxmlElement('w:r'); rp = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT_NAME)
            rp.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rp.append(sz)
            c = OxmlElement('w:color'); c.set(qn('w:val'), ORANGE_H); rp.append(c)
            u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rp.append(u)
            if bold: rp.append(OxmlElement('w:b'))
            if italic: rp.append(OxmlElement('w:i'))
            r.append(rp)
            t = OxmlElement('w:t'); t.text = label.upper() if caps else label; r.append(t)
            h.append(r); p._p.append(h)

        def emit(chunk, is_bold, is_code):
            if chunk == '': return
            r = p.add_run(chunk.upper() if caps else chunk); font(r)
            r.font.size = Pt(size); r.font.color.rgb = color
            r.bold = bold or is_bold; r.italic = italic or is_code
            if tracking is not None:
                rp = r._element.get_or_add_rPr(); sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(tracking)); rp.append(sp)
        def emit_bold(chunk):
            # `code` nests inside **bold** all the time, because the Problem headings are
            # written as **Problem N: `keyword` does X.** Without this second pass the
            # backticks printed literally inside the bolded run.
            at = 0
            for c in _CODE.finditer(chunk):
                emit(chunk[at:c.start()], True, False)
                emit(c.group(1), True, True)
                at = c.end()
            emit(chunk[at:], True, False)
        pos = 0
        for m in _INLINE.finditer(text):
            emit(text[pos:m.start()], False, False)
            if m.group(1) is not None: emit_bold(m.group(1))
            elif m.group(2) is not None: emit(m.group(2), False, True)
            else: emit_link(m.group(3), m.group(4))
            pos = m.end()
        emit(text[pos:], False, False)

    def para(text, **k):
        p = doc.add_paragraph(); runs(p, text, **k); return p

    def cantsplit(t):
        for row in t.rows:
            row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))

    def shade(c, h):
        s = OxmlElement('w:shd'); s.set(qn('w:val'), 'clear'); s.set(qn('w:fill'), h); c._tc.get_or_add_tcPr().append(s)

    def topborder(c, h, sz=24):
        b = OxmlElement('w:tcBorders'); t = OxmlElement('w:top')
        t.set(qn('w:val'), 'single'); t.set(qn('w:sz'), str(sz)); t.set(qn('w:color'), h); b.append(t)
        c._tc.get_or_add_tcPr().append(b)

    def no_table_borders(t):
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + e); x.set(qn('w:val'), 'nil'); b.append(x)
        t._tbl.tblPr.append(b)

    def no_cell_padding(t):
        # A table cell carries 0.08in of side padding by default, so a right-aligned run
        # in the last cell stops ~5.75pt short of the table edge. That is the same visible
        # error the tab-stop layout had, arriving by a different route, and it is only
        # findable by rendering the page and measuring the ink.
        m = OxmlElement('w:tblCellMar')
        for e in ('top', 'left', 'bottom', 'right'):
            x = OxmlElement('w:' + e); x.set(qn('w:w'), '0'); x.set(qn('w:type'), 'dxa'); m.append(x)
        t._tbl.tblPr.append(m)

    def col_widths(t, widths):
        # Widths are EMU, not inches, and must be split so they sum back to the total
        # exactly. Converting each column through Inches() truncates every one of them, and
        # a table that is a few EMU narrower than the text block is the whole class of bug
        # this layout exists to avoid.
        #
        # Setting cell widths alone is not enough either: python-docx lays a table out with
        # an equal tblGrid, and both Word and Google Docs size the columns from that grid
        # unless the table is explicitly fixed-layout. Without this the running header and
        # footer columns stayed equal thirds, which wrapped a long footer label onto a
        # second line.
        tl = OxmlElement('w:tblLayout'); tl.set(qn('w:type'), 'fixed'); t._tbl.tblPr.append(tl)
        for col, w in zip(t.columns, widths):
            col.width = Emu(int(w))
        for row in t.rows:
            for c, w in zip(row.cells, widths):
                c.width = Emu(int(w))

    def split_width(total, *shares):
        """Split `total` EMU by share so the columns sum back to it exactly.

        A tblGrid stores widths in TWIPS (1 twip = 635 EMU), so an arbitrary EMU value is
        rounded on write and three rounded columns no longer add up to the text width.
        Splitting in twips and giving the remainder to the last column keeps the total
        exact in the unit the file actually stores.
        """
        twips = int(total) // 635
        parts = [int(twips * s) for s in shares[:-1]]
        parts.append(twips - sum(parts))
        return [p * 635 for p in parts]

    def compact(p):
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

    def field(p, instruction, display):
        # Every run in the field carries the footer's own size/colour, not just the cached
        # display run. Word and Google Docs re-render the field using the field runs'
        # formatting, so an unstyled begin/instrText/separate/end leaks Normal (10.5pt Ink)
        # into the footer and the number renders bigger and darker than "page" / "of".
        def fld_run():
            r = p.add_run(); font(r); r.font.size = Pt(8); r.font.color.rgb = MIST
            return r
        r = fld_run(); begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin'); r._r.append(begin)
        r = fld_run(); instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = f' {instruction} '
        r._r.append(instr)
        r = fld_run(); separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate'); r._r.append(separate)
        r = p.add_run(display); font(r); r.font.size = Pt(8); r.font.color.rgb = MIST
        r = fld_run(); end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); r._r.append(end)

    def orange_rule():
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.keep_with_next = True; p.add_run()
        pPr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
        bo.set(qn('w:val'), 'single'); bo.set(qn('w:sz'), '24'); bo.set(qn('w:space'), '1'); bo.set(qn('w:color'), ORANGE_H)
        bd.append(bo); pPr.append(bd); p.paragraph_format.right_indent = Inches(5.4)

    def data_table(rows):
        t = doc.add_table(rows=0, cols=len(rows[0]))
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(len(rows[0])):
                txt = row[ci] if ci < len(row) else ''
                cells[ci].text = ''
                if ri == 0: shade(cells[ci], INK_H)
                elif ri % 2 == 0: shade(cells[ci], CLOUD_H)
                pc = cells[ci].paragraphs[0]; pc.paragraph_format.space_after = Pt(2); pc.paragraph_format.space_before = Pt(2)
                if ri == 0: pc.paragraph_format.keep_with_next = True
                runs(pc, txt, size=8.5, color=(WHITE if ri == 0 else INK), bold=(ri == 0))
        # Repeat the header row and keep it attached to the first data row. Google Docs
        # otherwise can strand a dark table header at the bottom of a page after import.
        trpr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trpr.append(th)
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'bottom', 'insideH'):
            x = OxmlElement('w:' + e); x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), '4'); x.set(qn('w:color'), MISTLINE_H); b.append(x)
        for e in ('left', 'right', 'insideV'):
            x = OxmlElement('w:' + e); x.set(qn('w:val'), 'none'); b.append(x)
        t._tbl.tblPr.append(b); cantsplit(t)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def kpi_cards(items):
        t = doc.add_table(rows=0, cols=len(items))
        nb = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + e); x.set(qn('w:val'), 'none'); nb.append(x)
        t._tbl.tblPr.append(nb)
        cells = t.add_row().cells
        for ci, (num, label, sub) in enumerate(items):
            c = cells[ci]; c.text = ''; shade(c, CLOUD_H); topborder(c, ORANGE_H, 24)
            p0 = c.paragraphs[0]; p0.paragraph_format.space_before = Pt(6); p0.paragraph_format.space_after = Pt(0)
            runs(p0, num, size=25, color=INK, bold=True)
            p1 = c.add_paragraph(); p1.paragraph_format.space_after = Pt(0)
            runs(p1, label, size=8, color=STEEL, bold=True, caps=True, tracking=20)
            p2 = c.add_paragraph(); p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(6)
            if sub: runs(p2, sub, size=8.5, color=ORANGE)
        cantsplit(t)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def image(path, caption):
        if not Path(path).exists():
            para(f"[missing image: {Path(path).name}]", size=8.5, color=STEEL, italic=True); return
        from docx.image.image import Image as DocxImage
        source = DocxImage.from_file(str(path))
        aspect = source.height / source.width if source.width else 1
        max_width = 5.0 if Path(path).name.startswith('fig_') else 5.6
        width_inches = min(max_width, 8.4 / aspect) if aspect else max_width
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width_inches))
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(10)
        runs(c, caption, size=8.5, color=STEEL, italic=True)

    def doc_title(text):
        """Page-one document title. Only drawn when there is no cover to carry it."""
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True; runs(p, text, size=24, color=INK, bold=True)

    def doc_byline(text):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True; runs(p, text, size=9.5, color=STEEL, bold=False)
        orange_rule()

    def h2(text):
        orange_rule()
        p = doc.add_paragraph(style='Heading 1'); p.paragraph_format.space_after = Pt(8); p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.keep_with_next = True; runs(p, text, size=18, color=INK, bold=True)

    def h3(text):
        p = doc.add_paragraph(style='Heading 2'); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True; runs(p, text, size=12.5, color=INK, bold=True)

    def h4(text):
        p = doc.add_paragraph(style='Heading 3'); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True; runs(p, text, size=11, color=ORANGE, bold=True)

    def lever(n, title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True; runs(p, f"LEVER {n}", size=9, color=ORANGE, bold=True, caps=True, tracking=30)
        p2 = doc.add_paragraph(style='Heading 2'); p2.paragraph_format.space_after = Pt(6); p2.paragraph_format.keep_with_next = True
        runs(p2, title, size=14, color=INK, bold=True)

    def note(text):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(9)
        pPr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
        lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '18'); lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), ORANGE_H)
        bd.append(lf); pPr.append(bd); p.paragraph_format.left_indent = Inches(0.12)
        runs(p, text, size=9.5, color=STEEL, italic=True)

    def bullet(text, num=False):
        p = doc.add_paragraph(style='List Number' if num else 'List Bullet'); runs(p, text)

    # Use one page style for the whole document. A separate zero-margin cover section caused
    # LibreOffice to alternate the cover and body styles on automatically paginated pages.
    # Keeping the cover inline also avoids first-page-style resets after automatic page breaks.
    # The small white frame is preferable to clipped headings or missing running furniture.
    body = doc.sections[0]
    body.page_width = A4_W; body.page_height = A4_H
    body.top_margin = Inches(0.85); body.bottom_margin = Inches(0.8)
    body.left_margin = Inches(0.85); body.right_margin = Inches(0.85)
    body.header_distance = Inches(0.25); body.footer_distance = Inches(0.25)
    # One source of truth for the text width. It was written out as the literal 6.57 in
    # three places and 3.285 in a fourth, and it has already drifted once (an older footer
    # used 6.55), which silently desynchronizes the running furniture from the body.
    content_w = body.page_width - body.left_margin - body.right_margin
    if cover_png:
        cp = doc.add_paragraph(); compact(cp)
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cr = cp.add_run(); cr.font.size = Pt(1)
        cr.add_picture(str(cover_png), width=content_w)
        doc.add_page_break()

    # Running furniture. Populate both right-page and left-page parts, while keeping the
    # first-page variant disabled. LibreOffice applies its paired page styles even when a
    # DOCX asks for shared furniture; duplicating the contract into both parts keeps every
    # page branded without triggering its fragile First Page overflow behavior.
    run = _running_text(cfg, M)
    logo = Path(brand_dir) / _B["assets"].get("logo_black", "logo_black.png")
    doc.settings.odd_and_even_pages_header_footer = True
    # An empty first-page header and footer, so the COVER carries no furniture in the DOCX
    # itself. Relying on the post-conversion useFirstPageHeaderFooter request meant a skipped
    # or failed normalization shipped a cover with a header stamped across it.
    # Only when a cover exists: with cover=False page one is ordinary content, and blanking
    # its furniture stripped the running header, the footer and the page number from it.
    body.different_first_page_header_footer = bool(cover_png)

    # FIXED-WIDTH TABLES, never tab stops.
    #
    # python-docx builds header and footer parts from a default part whose paragraph carries
    # the built-in Header/Footer styles, and those styles define US-LETTER tab stops: centre
    # 4680 twips (3.25in) and right 9360 twips (6.5in). This document is A4 with 0.85in
    # margins, so the content width is 6.57in and the centre is 3.285in. OOXML merges
    # paragraph tabs with style tabs additively and replaces a stop only at an IDENTICAL
    # position, so adding 3.285in and 6.57in does not clear 3.25in and 6.5in, it appends to
    # them. Traversal reaches the Letter stops first and the correct ones are never used,
    # which is why a delivered audit had its header label and footer URL sitting 5pt short of
    # the text edge directly below them, and the page counter 2.5pt left of centre. A table
    # cell width is absolute and inherits nothing, so it cannot drift that way.
    def populate_header(header, blank=False):
        header.is_linked_to_previous = False
        header.paragraphs[0].text = ''
        compact(header.paragraphs[0])
        if blank:
            return
        t = header.add_table(rows=1, cols=2, width=content_w)
        no_table_borders(t); no_cell_padding(t)
        col_widths(t, split_width(content_w, 0.5, 0.5))
        left, right = t.rows[0].cells
        lp = left.paragraphs[0]; compact(lp)
        if logo.exists():
            lp.add_run().add_picture(str(logo), width=Inches(0.9))
        elif _B.get("agency_name"):
            runs(lp, _B["agency_name"], size=8, color=INK, bold=True)
        rp = right.paragraphs[0]; compact(rp)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runs(rp, run["header_right"], size=7.5, color=MIST, bold=True, tracking=12)

    page_word, of_word = _footer_page_words(cfg)

    def populate_footer(footer, blank=False):
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = ''
        compact(footer.paragraphs[0])
        if blank:
            return
        t = footer.add_table(rows=1, cols=3, width=content_w)
        no_table_borders(t); no_cell_padding(t)
        # The left zone carries the longest string, so it gets the slack. Equal thirds is
        # what used to wrap a long client name onto a second line.
        col_widths(t, split_width(content_w, 0.42, 0.24, 0.34))
        left, mid, right = t.rows[0].cells
        lp = left.paragraphs[0]; compact(lp)
        runs(lp, run["footer_left"], size=8, color=MIST)
        mp = mid.paragraphs[0]; compact(mp)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runs(mp, page_word, size=8, color=MIST); field(mp, "PAGE", "1")
        runs(mp, of_word, size=8, color=MIST); field(mp, "NUMPAGES", "1")
        rp = right.paragraphs[0]; compact(rp)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runs(rp, run["footer_right"], size=8, color=MIST)

    populate_header(body.header)
    populate_header(body.even_page_header)
    if cover_png: populate_header(body.first_page_header, blank=True)
    populate_footer(body.footer)
    populate_footer(body.even_page_footer)
    if cover_png: populate_footer(body.first_page_footer, blank=True)

    # render blocks (+ KPI after the verdict/summary h2); -1 disables cards for non-audit docs
    kpi_idx = _kpi_after(blocks) if (M.get("custom_kpis") or M.get("totals")) else -1
    for i, (kind, payload) in enumerate(blocks):
        if kind == "h2":
            h2(payload)
            if i == kpi_idx: kpi_cards(_kpis(M))
        elif kind == "doctitle":
            if not cover_png: doc_title(payload)
        elif kind == "byline":
            if not cover_png: doc_byline(payload)
        elif kind == "h3": h3(payload)
        elif kind == "h4": h4(payload)
        elif kind == "lever": lever(*payload)
        elif kind == "p": para(payload)
        elif kind == "note": note(payload)
        elif kind == "bul": bullet(payload)
        elif kind == "num": bullet(payload, num=True)
        elif kind == "table": data_table(payload)
        elif kind == "img": image(*payload)
        elif kind == "pagebreak": doc.add_page_break()
    # NUMPAGES has no cached value we can compute here, so it ships as "1". Without this
    # flag Word and Google Docs render that stale cache and a multi-page audit reads
    # "page 1 of 1". updateFields makes them recompute the footer counters on open.
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)
    doc.save(out)
    return out


# ------------------------------ entry ------------------------------
def render(cfg, outdir, scaffold_md, cover=False, brand_dir=None):
    """Render the branded .docx from the narrative markdown. Returns dict of outputs.
    Raises on hard failure; callers should fall back to md_to_docx."""
    outdir = Path(outdir).resolve(); scaffold_md = Path(scaffold_md)
    _apply_branding(_branding.load_branding(cfg))
    brand_dir = Path(brand_dir or _bcfg(cfg, "brand_dir") or _B["assets"]["brand_dir"] or (HERE / "brand"))
    M = json.loads((outdir / "metrics.json").read_text())
    title, blocks = parse_markdown(scaffold_md.read_text(), scaffold_md.parent)
    stem = scaffold_md.name.replace("_SCAFFOLD.md", "").replace(".md", "")
    out_docx = outdir / f"{stem}_BRANDED.docx"
    result = {}

    cover_png = None
    if cover:
        need = [FONT_FILE, _B["assets"]["logo_white"]]
        if all((brand_dir / n).exists() for n in need):
            from brand_cover import build_cover
            cover_png = outdir / "_cover.png"
            build_cover(str(cover_png), **_cover_kwargs(cfg, M, blocks, brand_dir))
        else:
            print(f"[brand] cover skipped — assets missing in {brand_dir} (run prepare_brand_assets.py)")

    _render_docx(blocks, M, cfg, brand_dir, cover_png, out_docx)
    result["docx"] = out_docx; print("[brand] wrote", out_docx.name)
    return result


if __name__ == "__main__":
    import argparse
    from analyze_audit import load_config
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", required=True); ap.add_argument("--outdir", required=True)
    ap.add_argument("--md", required=True); ap.add_argument("--cover", action="store_true")
    a = ap.parse_args()
    render(load_config(a.config), a.outdir, a.md, cover=a.cover)
