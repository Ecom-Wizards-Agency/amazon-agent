#!/usr/bin/env python3
"""Minimal Markdown -> .docx converter for the audit narrative.
Handles: #/##/### headings, **bold** inline, '- ' bullets, '1.' numbered, and | pipe tables.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(sys.argv[1]); OUT = Path(sys.argv[2])

def _rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))
try:  # palette from the agency branding file; hard fallback — this script must never fail
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from branding import load_branding
    _pal = load_branding({})["palette_xlsx"]
    INK = _rgb(_pal["ink"]); CORAL = _rgb(_pal["coral"])
except Exception:
    INK = RGBColor(0x1E,0x24,0x2C); CORAL = RGBColor(0x3B,0x6E,0xF6)

doc = Document()
style = doc.styles['Normal']; style.font.name='Calibri'; style.font.size=Pt(10.5)

def add_runs(p, text):
    # split on **bold**
    for i,part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if part=='': continue
        run=p.add_run(part); run.bold = (i%2==1)
        run.font.color.rgb = INK

def flush_table(rows):
    if not rows: return
    ncol=max(len(r) for r in rows)
    t=doc.add_table(rows=0, cols=ncol); t.style='Light Grid Accent 1'
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci in range(ncol):
            txt=row[ci] if ci<len(row) else ''
            cell=cells[ci]; cell.text=''
            p=cell.paragraphs[0]
            for j,part in enumerate(re.split(r'\*\*(.+?)\*\*', txt)):
                if part=='': continue
                run=p.add_run(part); run.bold=(j%2==1) or (ri==0)
                run.font.size=Pt(9)
    doc.add_paragraph()

lines = SRC.read_text(encoding='utf-8').splitlines()
tbl=[]
for ln in lines:
    s=ln.rstrip()
    if re.match(r'^\|.*\|\s*$', s):
        cells=[c.strip() for c in s.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c): continue  # separator
        tbl.append(cells); continue
    else:
        if tbl: flush_table(tbl); tbl=[]
    if not s.strip(): continue
    if s.startswith('### '): doc.add_heading(s[4:],level=3)
    elif s.startswith('## '): doc.add_heading(s[3:],level=2)
    elif s.startswith('# '):
        h=doc.add_heading('',level=1); add_runs(h,s[2:])
    elif re.match(r'^\s*[-*] ', s):
        p=doc.add_paragraph(style='List Bullet'); add_runs(p, re.sub(r'^\s*[-*] ','',s))
    elif re.match(r'^\s*\d+\. ', s):
        p=doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\s*\d+\. ','',s))
    else:
        p=doc.add_paragraph(); add_runs(p,s)
if tbl: flush_table(tbl)
doc.save(OUT)
print("Saved", OUT.name)
