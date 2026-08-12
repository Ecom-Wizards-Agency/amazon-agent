"""The running header and footer are laid out by fixed-width tables, at the text width.

They used to be tab-based, and shipped a client audit whose header label and footer URL
sat 5pt short of the body text edge below them, with the page counter 2.5pt left of
centre. The cause was not visible in any test: python-docx's built-in Header/Footer
styles carry US-LETTER tab stops (centre 3.25in, right 6.5in), OOXML merges paragraph
tabs with style tabs additively rather than replacing them, and the renderer's correct
A4 stops at 3.285in and 6.57in were therefore appended after the Letter ones and never
reached.

So these tests assert POSITIONS and structure, never tab characters. A test that checks
for "\\t\\t" in the header text passes while the header is still 5pt short.
"""
import json
import sys
import tempfile
import unittest
from pathlib import Path

TOOL_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(TOOL_DIR))

import native_doc_normalize as N  # noqa: E402
import render_branded  # noqa: E402

EMU_PER_INCH = 914400
METRICS = {
    "client": "Fixture", "currency": "USD", "marketplaces": ["US"], "breakeven": 0.5,
    "channels_present": ["SP"],
    "windows": {"ads": "2026-08-01..2026-08-28", "business_report": "same", "sqp_weeks": []},
    "totals": {"spend": 100.0, "sales": 250.0, "br_total_sales": 500.0,
               "acos": 0.4, "roas": 2.5, "tacos": 0.2, "ad_dependency": 0.5},
    "searchterm_bucket": {"Branded": {"spend": 40.0, "sales": 120.0, "acos": 1/3, "cvr": 0.2}},
    "placement": {}, "business_report": {"rows": []},
}


def render(**cfg_extra):
    tmp = tempfile.TemporaryDirectory()
    root = Path(tmp.name)
    cfg = {"client": "Fixture", "marketplaces": ["US"], "date": "2026-08-12", **cfg_extra}
    (root / "metrics.json").write_text(json.dumps(METRICS))
    md = root / "Fixture_US_Sales_Audit_SCAFFOLD.md"
    md.write_text("# Fixture\n\n## Ads Summary\n\nA line of body copy.\n")
    out = render_branded.render(cfg, root, md, cover=False)
    from docx import Document
    doc = Document(str(out["docx"] if isinstance(out, dict) else out))
    tmp_ref = tmp  # keep the directory alive for the caller
    return doc, tmp_ref


class RunningFurnitureTest(unittest.TestCase):
    def setUp(self):
        self.doc, self._tmp = render()
        self.section = self.doc.sections[0]
        self.content_w = (self.section.page_width
                          - self.section.left_margin
                          - self.section.right_margin)

    def tearDown(self):
        self._tmp.cleanup()

    def test_header_is_a_two_cell_table_at_the_text_width(self):
        header = self.section.header
        self.assertEqual(1, len(header.tables), "header must be laid out by one table")
        t = header.tables[0]
        self.assertEqual(2, len(t.columns))
        self.assertEqual(self.content_w, sum(c.width for c in t.columns),
                         "header table must span exactly the text width")

    def test_footer_is_a_three_cell_table_at_the_text_width(self):
        footer = self.section.footer
        self.assertEqual(1, len(footer.tables))
        t = footer.tables[0]
        self.assertEqual(3, len(t.columns))
        self.assertEqual(self.content_w, sum(c.width for c in t.columns),
                         "footer table must span exactly the text width")

    def test_no_tab_stops_or_tab_characters_survive_anywhere(self):
        """The whole defect was tabs. If one comes back, the Letter stops come back with it."""
        for part, name in ((self.section.header, "header"), (self.section.footer, "footer")):
            xml = part.part.element.xml
            self.assertNotIn("<w:tab/>", xml, f"{name} must contain no tab characters")
            self.assertNotIn("<w:tabs>", xml, f"{name} must define no tab stops")

    def test_footer_zones_are_left_centre_right(self):
        cells = self.section.footer.tables[0].rows[0].cells
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.assertIn("Fixture", cells[0].text)
        self.assertEqual(WD_ALIGN_PARAGRAPH.CENTER, cells[1].paragraphs[0].alignment)
        self.assertEqual(WD_ALIGN_PARAGRAPH.RIGHT, cells[2].paragraphs[0].alignment)
        xml = self.section.footer.part.element.xml
        self.assertIn("PAGE", xml)
        self.assertIn("NUMPAGES", xml)

    def test_cells_have_no_padding(self):
        """0.08in of default cell padding stops a right-aligned run short of the edge,
        which is the same visible error the old tab layout had."""
        for part, name in ((self.section.header, "header"), (self.section.footer, "footer")):
            xml = part.tables[0]._tbl.xml
            self.assertIn("tblCellMar", xml, f"{name} table must zero its cell padding")

    def test_header_label_is_right_aligned(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        right = self.section.header.tables[0].rows[0].cells[1]
        self.assertEqual(WD_ALIGN_PARAGRAPH.RIGHT, right.paragraphs[0].alignment)
        self.assertIn("ACCOUNT AUDIT", right.text.upper())

    def test_the_cover_page_carries_no_furniture(self):
        """A skipped normalization must not be able to stamp a header over the cover."""
        self.assertTrue(self.section.different_first_page_header_footer)
        self.assertEqual(0, len(self.section.first_page_header.tables))
        self.assertEqual(0, len(self.section.first_page_footer.tables))
        self.assertEqual("", self.section.first_page_header.paragraphs[0].text.strip())

    def test_doc_label_default_and_override(self):
        self.assertEqual("Account Audit", render_branded._doc_label({}))
        self.assertEqual("Amazon Account Review",
                         render_branded._doc_label({"branding": {"doc_label": "Amazon Account Review"}}))


class NormalizerTest(unittest.TestCase):
    """The normalizer reads the real API shape and repairs only the cover."""

    def _document(self, nested):
        page = {"width": {"magnitude": 595.28, "unit": "PT"},
                "height": {"magnitude": 841.89, "unit": "PT"}}
        doc_tab = {
            "body": {"content": [
                {"startIndex": 0, "endIndex": 1, "sectionBreak": {}},
                {"startIndex": 1, "endIndex": 3, "paragraph": {"elements": [
                    {"startIndex": 1, "endIndex": 2,
                     "inlineObjectElement": {"inlineObjectId": "i.0"}}]}},
                {"startIndex": 3, "endIndex": 5, "paragraph": {"elements": [
                    {"pageBreak": {}}]}},
            ]},
            "documentStyle": {"pageSize": page},
            "inlineObjects": {"i.0": {"inlineObjectProperties": {"embeddedObject": {
                "imageProperties": {"contentUri": "https://example/cover.png"}}}}},
        }
        if nested:
            return {"tabs": [{"documentTab": doc_tab, "tabProperties": {"tabId": "t.0"}}]}
        return {"tabs": [{**doc_tab, "tabId": "t.0"}]}

    def test_reads_the_real_nested_api_shape(self):
        """documents.get nests content under documentTab; reading only the flat shape
        made every live response look like a document with no cover."""
        for nested in (True, False):
            reqs = N.build_native_account_audit_requests(self._document(nested))
            kinds = [list(r)[0] for r in reqs]
            self.assertEqual(["deleteContentRange", "insertSectionBreak", "updateSectionStyle",
                              "deleteContentRange", "insertInlineImage",
                              "updateTextStyle", "updateParagraphStyle",
                              "updateDocumentStyle"], kinds)

    def test_cover_fills_the_page_apart_from_one_hairline(self):
        """Docs inserts a paragraph before a section break and it sits in the cover's
        zero-margin section, so a cover of the full page height spills it onto a blank
        second page. The cover gives back exactly that hairline and no more."""
        reqs = N.build_native_account_audit_requests(self._document(True))
        img = next(r["insertInlineImage"] for r in reqs if "insertInlineImage" in r)
        self.assertEqual(595.28, img["objectSize"]["width"]["magnitude"])
        self.assertEqual(841.89 - N.TRAILING_LINE, img["objectSize"]["height"]["magnitude"])
        self.assertLessEqual(N.TRAILING_LINE, 2.0, "the sliver must stay invisible")

    def test_the_section_break_paragraph_is_collapsed(self):
        reqs = N.build_native_account_audit_requests(self._document(True))
        size = next(r["updateTextStyle"] for r in reqs if "updateTextStyle" in r)
        self.assertEqual(1, size["textStyle"]["fontSize"]["magnitude"])
        para = next(r["updateParagraphStyle"] for r in reqs if "updateParagraphStyle" in r)
        self.assertEqual(100, para["paragraphStyle"]["lineSpacing"])

    def test_it_never_touches_header_alignment(self):
        reqs = N.build_native_account_audit_requests(self._document(True))
        self.assertNotIn("insertText", [list(r)[0] for r in reqs])

    def test_zero_margin_accepts_an_omitted_magnitude(self):
        """The API drops `magnitude` for a zero value, so `== 0` fails on a correct doc."""
        self.assertTrue(N.zero_margin({k: {"unit": "PT"} for k in
                                       ("marginTop", "marginBottom", "marginLeft", "marginRight")}))
        self.assertFalse(N.zero_margin({"marginTop": {"magnitude": 61.2, "unit": "PT"}}))

    def test_has_cover_is_false_for_a_monthly_audit(self):
        self.assertFalse(N.has_cover({"tabs": [{"documentTab": {"body": {"content": []}},
                                                "tabProperties": {"tabId": "t.0"}}]}))


if __name__ == "__main__":
    unittest.main(verbosity=2)
