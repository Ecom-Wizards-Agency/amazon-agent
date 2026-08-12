import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


TOOL_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(TOOL_DIR))

import render_branded  # noqa: E402


class PriorityHeadingBrandingTest(unittest.TestCase):
    def test_heading_two_style_and_priority_run_use_ink(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "audit.docx"
            render_branded._render_docx(
                [
                    ("h2", "Problems and Solutions"),
                    ("h3", "Priority 1: Test the highest-impact fix."),
                    ("p", "The evidence supports a focused test. I would measure the result."),
                ],
                {},
                {"client": "Fixture", "marketplaces": ["US"]},
                root,
                None,
                output,
            )

            doc = Document(output)
            heading_two = doc.styles["Heading 2"]
            self.assertEqual(render_branded.FONT_NAME, heading_two.font.name)
            self.assertEqual(12.5, heading_two.font.size.pt)
            self.assertEqual(render_branded.INK_H, str(heading_two.font.color.rgb))
            self.assertTrue(heading_two.font.bold)

            priority = next(p for p in doc.paragraphs if p.text.startswith("Priority 1:"))
            self.assertEqual("Heading 2", priority.style.name)
            self.assertEqual(render_branded.INK_H, str(priority.runs[0].font.color.rgb))
            self.assertTrue(priority.runs[0].bold)

            support = next(p for p in doc.paragraphs if p.text.startswith("The evidence"))
            self.assertEqual("Normal", support.style.name)
            self.assertEqual(render_branded.INK_H, str(support.runs[0].font.color.rgb))
            self.assertFalse(support.runs[0].bold)
            self.assertFalse(support.runs[0].italic)


if __name__ == "__main__":
    unittest.main()
