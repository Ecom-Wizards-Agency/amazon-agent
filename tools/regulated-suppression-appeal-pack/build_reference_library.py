#!/usr/bin/env python3
"""Build clean DOCX templates and a fictional suppression-appeal reference pack."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "203746"
TEAL = "167A78"
PALE = "EAF2F2"
LIGHT = "F2F4F5"
WHITE = "FFFFFF"
GRAY = "5D666D"

MODULES = (
    ("01-appeal", "Appeal and Request for Manual Review"),
    ("02-technical-evidence-and-claim-register", "Technical Evidence and Claim Register"),
    ("03-manufacturer-declaration", "Manufacturer, Product, and Component Declaration"),
    ("04-catalog-correction-and-processing-proof", "Catalog Correction and Processing Proof"),
    ("05-compliance-sop", "Regulated Product Content Compliance SOP"),
    ("06-training-guide", "Amazon Content Compliance Training Guide"),
    ("07-manual-review", "Instruction Manual Compliance Review"),
    ("08-packaging-and-label-review", "Packaging and Label Compliance Review"),
    ("09-evidence-index", "Evidence Index"),
    ("10-submission-checklist", "Submission Readiness Checklist"),
)


@dataclass(frozen=True)
class CaseData:
    mode: str

    def value(self, key: str) -> str:
        template = {
            "brand": "[BRAND NAME]",
            "seller": "[SELLER LEGAL NAME]",
            "marketplace": "[MARKETPLACE]",
            "case_id": "[CASE ID]",
            "notice_date": "[AMAZON NOTICE DATE]",
            "issue": "[SUPPRESSION ISSUE]",
            "product": "[PRODUCT FAMILY]",
            "asin_1": "[ASIN 1]",
            "asin_2": "[ASIN 2]",
            "sku_1": "[SKU 1]",
            "sku_2": "[SKU 2]",
            "manufacturer": "[MANUFACTURER LEGAL NAME]",
            "component_manufacturer": "[COMPONENT MANUFACTURER]",
            "model": "[MODEL OR COMPONENT]",
            "report": "[TEST REPORT NUMBER]",
            "laboratory": "[INDEPENDENT LABORATORY]",
            "specification": "[TECHNICAL SPECIFICATION]",
            "threshold": "[NOTICE-SPECIFIC COMPARISON]",
            "batch": "[AMAZON BATCH ID]",
            "owner": "[COMPLIANCE OWNER]",
            "effective_date": "[EFFECTIVE DATE]",
            "manual_version": "[MANUAL VERSION]",
            "packaging_version": "[PACKAGING VERSION]",
        }
        example = {
            "brand": "Aster & Vale",
            "seller": "Example Commerce LLC",
            "marketplace": "Amazon.com US",
            "case_id": "EXAMPLE-CASE-2401",
            "notice_date": "January 15, 2026",
            "issue": "Fictional restricted-product review involving cosmetic and device-positioning claims",
            "product": "Aster & Vale Surface Applicator Kit",
            "asin_1": "EXAMPLE-ASIN-01",
            "asin_2": "EXAMPLE-ASIN-02",
            "sku_1": "EXAMPLE-SKU-01",
            "sku_2": "EXAMPLE-SKU-02",
            "manufacturer": "Example Contract Manufacturing Ltd.",
            "component_manufacturer": "Example Components Ltd.",
            "model": "AV-SURFACE-12",
            "report": "EXAMPLE-TR-1042",
            "laboratory": "Example Independent Laboratory",
            "specification": "Fictional measured contact-point exposure of 0.12 mm",
            "threshold": "0.3 mm comparison stated in the fictional Amazon notice",
            "batch": "EXAMPLE-BATCH-8102",
            "owner": "Example Compliance Lead",
            "effective_date": "February 1, 2026",
            "manual_version": "EXAMPLE-IFU-2.0",
            "packaging_version": "EXAMPLE-PKG-2.0",
        }
        return (template if self.mode == "template" else example)[key]

    @property
    def label(self) -> str:
        return "Reusable Template" if self.mode == "template" else "Fictional Worked Example"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def configure_document(doc: Document, data: CaseData) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.54)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, color, before, after in (
        ("Title", 20, NAVY, 0, 6),
        ("Heading 1", 13.5, TEAL, 9, 3),
        ("Heading 2", 10.5, NAVY, 7, 2),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "REGULATED PRODUCT SUPPRESSION APPEAL SYSTEM"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.runs[0]
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    footer = section.footer.paragraphs[0]
    footer.text = (
        "Reusable template. Replace every placeholder before use."
        if data.mode == "template"
        else "Fictional training example. Not legal advice or regulatory precedent."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.runs[0]
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = ""
    doc.core_properties.subject = ""
    doc.core_properties.keywords = ""
    doc.core_properties.comments = ""


def add_title(doc: Document, title: str, module: str, data: CaseData) -> None:
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run(data.label.upper())
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    heading = doc.add_paragraph(title, style="Title")
    heading.paragraph_format.keep_with_next = True

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(7)
    run = note.add_run(
        f"{module} | {data.value('product')} | {data.value('marketplace')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    add_metadata_table(
        doc,
        [
            ("Brand", data.value("brand")),
            ("Seller", data.value("seller")),
            ("Case", data.value("case_id")),
            ("Issue", data.value("issue")),
            ("Document status", "DRAFT FOR EVIDENCE REVIEW"),
            ("Final approval", "Victor Uhl, Master of Troubleshooting"),
        ],
    )


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.8)
    for index, (label, value) in enumerate(rows):
        left, right = table.rows[index].cells
        left.width = Inches(1.45)
        right.width = Inches(5.8)
        set_cell_shading(left, NAVY)
        set_cell_shading(right, LIGHT if index % 2 == 0 else WHITE)
        for cell in (left, right):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.font.size = Pt(8.2)
        r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section(doc: Document, title: str, paragraphs: list[str] = (), bullets: list[str] = ()) -> None:
    doc.add_heading(title, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    for text in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.paragraph_format.first_line_indent = Inches(-0.16)
        paragraph.add_run(text)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    set_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)


def add_grid(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.font.bold = True
        r.font.size = Pt(7.6)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(7.3)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_appeal(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Request",
        [
            f"We request a senior manual review of {d.value('asin_1')} and {d.value('asin_2')} "
            f"under Case {d.value('case_id')}. The affected products are variations of "
            f"{d.value('product')}.",
            "This submission addresses the exact concern in Amazon's notice, documents completed "
            "catalog and content corrections, and maps every material statement to supporting evidence.",
        ],
    )
    add_section(
        doc,
        "1. Root Cause",
        [
            "The prior review process considered individual phrases and catalog fields separately. "
            "It did not reconcile the combined product presentation across all variations and surfaces.",
            "The process also lacked a single owner with authority to stop publication when evidence, "
            "wording, or variation data conflicted.",
        ],
    )
    add_section(
        doc,
        "2. Completed Corrective Actions",
        bullets=[
            "Reviewed every affected ASIN, SKU, variation, and seller-controlled catalog surface.",
            "Removed or corrected the notice-relevant wording and unsupported implications.",
            f"Submitted the controlled catalog correction under {d.value('batch')}.",
            "Removed or corrected conflicting media and linked content where applicable.",
            "Verified packaging, instructions, quantities, and product identifiers against source records.",
        ],
    )
    add_section(
        doc,
        "3. Preventative Measures",
        bullets=[
            f"Assigned final review authority to {d.value('owner')}.",
            "Implemented product-accuracy review followed by regulated-content review.",
            "Added a publication stop rule for unsupported, conflicting, or incomplete content.",
            "Required post-processing catalog verification and evidence retention.",
            "Completed team training for the affected product class and content surfaces.",
        ],
    )
    add_section(
        doc,
        "4. Technical And Documentary Support",
        [
            f"The claim register maps the relevant product and component facts to {d.value('report')}, "
            f"issued by {d.value('laboratory')}, and to the signed manufacturer declaration.",
            f"The documented specification is: {d.value('specification')}. The comparison used in this "
            f"appeal is limited to {d.value('threshold')}. This statement is not extended beyond the "
            "scope of the cited evidence.",
        ],
    )
    add_callout(
        doc,
        "Requested Outcome",
        "Complete a substantive manual review of the attached evidence and reinstate the affected "
        "listings if Amazon confirms that the documented corrections and product facts satisfy the "
        "applicable review path.",
    )


def build_evidence(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Evidence Standard",
        [
            "Every material appeal statement must point to a named attachment and exact locator. "
            "A conclusion remains blocked when the evidence does not identify the tested product, "
            "responsible party, method, date, or affected Amazon record.",
        ],
    )
    add_grid(
        doc,
        ["Claim ID", "Proposed statement", "Evidence", "Locator", "Status"],
        [
            [
                "C-001",
                f"{d.value('model')} is used in both affected variations.",
                "Manufacturer declaration",
                "Section 4, product mapping",
                "Verified" if d.mode == "example" else "[STATUS]",
            ],
            [
                "C-002",
                d.value("specification"),
                d.value("report"),
                "Results table and conclusion",
                "Verified" if d.mode == "example" else "[STATUS]",
            ],
            [
                "C-003",
                f"Catalog correction reached the {d.value('batch')} processing stage.",
                "Amazon processing summary",
                "SKU-level result rows",
                "Verified" if d.mode == "example" else "[STATUS]",
            ],
        ],
        [0.62, 2.35, 1.35, 1.45, 0.75],
    )
    add_section(
        doc,
        "Technical Evidence Review",
        [
            f"Report reference: {d.value('report')}",
            f"Independent source: {d.value('laboratory')}",
            f"Tested model or component: {d.value('model')}",
            f"Recorded result: {d.value('specification')}",
        ],
        bullets=[
            "Confirm the report identifies the tested sample and method.",
            "Confirm the report is complete, legible, signed or authenticated as applicable.",
            "Confirm the manufacturer declaration maps the tested item to every affected variation.",
            "Record limitations and do not generalize beyond the tested sample.",
        ],
    )
    add_callout(
        doc,
        "Evidence Gap Rule",
        "If any material statement lacks a source and locator, remove the statement or keep the pack "
        "at Evidence incomplete.",
    )


def build_declaration(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Declaration",
        [
            f"This declaration is issued by {d.value('manufacturer')} concerning products supplied "
            f"to {d.value('seller')} under the {d.value('brand')} brand.",
        ],
    )
    add_section(
        doc,
        "1. Company And Manufacturing Relationships",
        bullets=[
            f"Brand and finished-product owner: {d.value('seller')}",
            f"Finished-product manufacturer or supplier: {d.value('manufacturer')}",
            f"Component manufacturer: {d.value('component_manufacturer')}",
            "The exact legal relationship must match company and transaction records.",
        ],
    )
    add_section(
        doc,
        "2. Product And Component Specification",
        [
            f"The relevant component is {d.value('model')}. The supported specification is "
            f"{d.value('specification')}.",
            "This confirmation is limited to facts supported by manufacturing, purchasing, production, "
            "and quality records held by the declarant.",
        ],
    )
    add_section(
        doc,
        "3. Independent Evidence",
        [
            f"The cited independent record is {d.value('report')} from {d.value('laboratory')}. "
            "The declarant confirms whether the tested sample is the same model incorporated into the "
            "finished products listed below.",
        ],
    )
    add_grid(
        doc,
        ["Amazon reference", "Seller SKU", "Variation", "Component", "Identifier"],
        [
            [d.value("asin_1"), d.value("sku_1"), "Variation 1", d.value("model"), "[IDENTIFIER]" if d.mode == "template" else "EXAMPLE-ID-01"],
            [d.value("asin_2"), d.value("sku_2"), "Variation 2", d.value("model"), "[IDENTIFIER]" if d.mode == "template" else "EXAMPLE-ID-02"],
        ],
        [1.35, 1.35, 1.2, 1.65, 1.6],
    )
    add_section(
        doc,
        "4. Authorized Confirmation",
        [
            "The undersigned confirms that the statements above are accurate according to the "
            "declarant's records and that supporting records can be produced on request.",
        ],
    )
    add_grid(
        doc,
        ["Authorized signatory", "Title", "Date", "Signature or stamp"],
        [[
            "[NAME]" if d.mode == "template" else "Example Authorized Signatory",
            "[TITLE]" if d.mode == "template" else "Quality Director",
            "[DATE]" if d.mode == "template" else "February 1, 2026",
            "[SIGNATURE OR STAMP]" if d.mode == "template" else "Required before real submission",
        ]],
        [2.0, 1.7, 1.3, 2.1],
    )


def build_catalog(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Catalog Correction Scope",
        [
            "This record distinguishes prepared, uploaded, processed, and verified states. A file is "
            "not described as successful until Amazon's processing summary confirms the SKU-level result.",
        ],
    )
    add_grid(
        doc,
        ["ASIN", "SKU", "Surface or field", "Prior issue", "Corrected standard"],
        [
            [d.value("asin_1"), d.value("sku_1"), "Title and bullets", "Notice-relevant claim", "Factual identity and contents"],
            [d.value("asin_1"), d.value("sku_1"), "Backend attributes", "Unsupported benefit language", "Neutral product facts"],
            [d.value("asin_2"), d.value("sku_2"), "Images and linked media", "Conflicting legacy content", "Reviewed or removed"],
        ],
        [1.15, 1.25, 1.35, 1.7, 1.7],
    )
    add_section(
        doc,
        "Processing Proof",
        bullets=[
            f"Batch or submission ID: {d.value('batch')}",
            "Upload file contains only intended SKUs and fields.",
            "Processing summary is attached and reviewed at SKU level.",
            "Warnings and rejected records are disclosed and resolved.",
        ],
    )
    add_section(
        doc,
        "Post-Processing Verification",
        bullets=[
            "Verify the selected seller account and marketplace.",
            "Verify frontend content for every affected variation.",
            "Verify backend values in a fresh report or edit view.",
            "Verify that removed A+, videos, images, and linked content did not return.",
            "Record screenshots, timestamps, reviewer, and unresolved catalog behavior.",
        ],
    )
    add_callout(
        doc,
        "State At Submission",
        "Prepared, uploaded, processed, and verified are separate facts. Select only the highest state "
        "supported by attached Amazon records.",
    )


def build_sop(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "1. Purpose And Scope",
        [
            f"This SOP governs Amazon content and supporting records for {d.value('product')} and any "
            "related regulated-product listing handled by the business.",
        ],
    )
    add_section(
        doc,
        "2. Roles",
        bullets=[
            "Product owner confirms identity, composition, quantity, specifications, and source records.",
            "Content owner prepares proposed catalog and media changes.",
            f"{d.value('owner')} performs the regulated-content review and controls publication.",
            "Victor provides final agency troubleshooting signoff for serious suppression appeal packs.",
        ],
    )
    add_section(
        doc,
        "3. Required Review Surfaces",
        bullets=[
            "Titles, bullets, descriptions, Item Highlights, images, A+ Content, and videos",
            "Search terms, intended use, directions, benefits, and other backend attributes",
            "Packaging, labels, manuals, inserts, and QR-linked content",
            "Every ASIN, SKU, variation, and seller contribution in the affected family",
        ],
    )
    add_section(
        doc,
        "4. Two-Stage Approval",
        [
            "Stage One verifies product accuracy against source records. Stage Two reviews regulated, "
            "restricted, therapeutic, disease, device, drug, supplement, and unsupported implications.",
            "Nothing may be published until both stages are complete.",
        ],
    )
    add_section(
        doc,
        "5. Publication Stop Rule",
        bullets=[
            "Stop when a claim is unsupported or ambiguous.",
            "Stop when variations or content surfaces conflict.",
            "Stop when required manufacturer, laboratory, or regulatory evidence is missing.",
            "Stop when a catalog correction lacks a successful processing result.",
            "Stop when final packaging, manuals, or linked content have not been reviewed.",
        ],
    )
    add_section(
        doc,
        "6. Submission And Verification Controls",
        bullets=[
            "Use controlled files containing only intended SKUs and fields.",
            "Retain the upload, processing summary, warnings, and SKU-level results.",
            "Verify frontend and backend results after processing.",
            "Reopen the case record when legacy content returns.",
        ],
    )
    add_section(
        doc,
        "7. Records, Training, And Review",
        [
            f"This SOP takes effect on {d.value('effective_date')}. Retain evidence, approvals, "
            "training records, catalog files, processing summaries, and post-publication checks with "
            "the case pack.",
        ],
    )


def build_training(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Learning Objectives",
        bullets=[
            "Recognize when the full product presentation creates a regulated or unsupported implication.",
            "Separate factual product information from marketing or therapeutic claims.",
            "Review the complete product family and every content surface.",
            "Distinguish prepared, uploaded, processed, and verified catalog states.",
            "Apply the stop and escalation rule.",
        ],
    )
    add_section(
        doc,
        "1. Main Rule",
        [
            "Amazon may evaluate the complete product presentation. A neutral title does not cure a "
            "claim that remains in an image, backend field, package, manual, video, or linked page.",
        ],
    )
    add_section(
        doc,
        "2. Permitted Content",
        bullets=[
            "Accurate product identity, contents, quantity, and ingredients",
            "Supported specifications and neutral mechanical operation",
            "Factual handling, storage, directions, and safety information",
            "Variation differences supported by product records",
        ],
    )
    add_section(
        doc,
        "3. Reject Or Escalate",
        bullets=[
            "Disease, treatment, prevention, cure, or therapeutic claims",
            "Unsupported body, skin, performance, absorption, delivery, or structure claims",
            "Clearance, registration, certification, testing, or endorsement statements without evidence",
            "Conflicting wording between listing, packaging, manual, and linked content",
        ],
    )
    add_section(
        doc,
        "4. Practical Scenarios",
        [
            "Scenario A: A visible claim is removed, but the same phrase remains in a backend benefit "
            "field. Action: stop and reconcile the complete content set.",
            "Scenario B: A file was uploaded, but no processing summary is available. Action: record it "
            "only as uploaded.",
            "Scenario C: Packaging links to an older manual. Action: block publication until the linked "
            "content is corrected and verified.",
        ],
    )
    add_section(
        doc,
        "5. Knowledge Check",
        bullets=[
            "Which surfaces must be reviewed together?",
            "What evidence supports each material product claim?",
            "When must publication stop?",
            "What proves that a catalog correction was processed and verified?",
            "Who has authority to approve the final appeal pack?",
        ],
    )


def build_manual(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Review Scope",
        [
            f"Manual reviewed: {d.value('manual_version')}. Review the complete instruction file, "
            "including cover, directions, expected-results language, warnings, endorsements, "
            "cross-sells, QR destinations, and images.",
        ],
    )
    add_grid(
        doc,
        ["Location", "Current wording or issue", "Risk", "Required action", "Evidence"],
        [
            ["Expected results", "[EXACT WORDING]" if d.mode == "template" else "Skin-result promise", "High", "Remove or rewrite neutrally", "Notice and claim standard"],
            ["Directions", "[EXACT WORDING]" if d.mode == "template" else "Application instruction", "Medium", "Keep factual and supported", "Approved operating record"],
            ["Safety", "[EXACT WORDING]" if d.mode == "template" else "Factual precaution", "Review", "Retain only necessary facts", "Safety record"],
            ["QR link", "[DESTINATION]" if d.mode == "template" else "Example instruction page", "High", "Verify final linked content", "Captured destination"],
        ],
        [1.1, 1.8, 0.65, 1.75, 1.55],
    )
    add_section(
        doc,
        "Required Checks",
        bullets=[
            "Remove unsupported results, absorption, delivery, treatment, disease, or therapeutic implications.",
            "Keep genuine safety instructions factual and proportionate.",
            "Remove endorsements or testing claims without exact-product evidence.",
            "Confirm terminology matches the listing and packaging.",
            "Confirm every referenced diagram, QR code, and linked page exists and is reviewed.",
            "Correct grammar and return one professionally formatted final version.",
        ],
    )
    add_callout(
        doc,
        "Disposition",
        "Do not mark the manual approved until every high-risk item is resolved and the final version "
        "is re-reviewed against the listing and packaging.",
    )


def build_packaging(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Review Scope",
        [
            f"Artwork or physical packaging reviewed: {d.value('packaging_version')}. Review each "
            "variation separately and capture all sides of the final consumer-facing package.",
        ],
    )
    add_grid(
        doc,
        ["Panel", "Claim or label issue", "Risk", "Required correction", "Verification"],
        [
            ["Front", "[WORDING]" if d.mode == "template" else "Unsupported result slogan", "High", "Remove or substantiate", "Final artwork and photo"],
            ["Back", "[WORDING]" if d.mode == "template" else "Instructions and ingredient block", "Review", "Match approved records", "Line-by-line comparison"],
            ["Side", "[WORDING]" if d.mode == "template" else "Certification statement", "High", "Clarify or remove", "Underlying certificate"],
            ["QR panel", "[DESTINATION]" if d.mode == "template" else "Example instruction page", "High", "Point to approved content only", "Live destination capture"],
        ],
        [0.8, 1.8, 0.65, 1.8, 1.65],
    )
    add_section(
        doc,
        "Label And Quantity Controls",
        bullets=[
            "Confirm ingredients match the component label, listing, and backend.",
            "Confirm unit count, net contents, serving or use quantity, and variation name.",
            "Label manufacturing, expiration, lot, and batch information clearly where applicable.",
            "Use certification, testing, facility, or endorsement claims only with exact supporting evidence.",
            "Remove production notes and dimensions from customer-facing evidence artwork.",
        ],
    )
    add_section(
        doc,
        "Final Evidence",
        bullets=[
            "Approved flat artwork for each variation",
            "Clear physical photos of front, back, left, right, top, and bottom",
            "Readable close-ups of ingredients, warnings, identifiers, and codes",
            "Verified QR destination and captured linked content",
            "Reviewer, date, version, and final disposition",
        ],
    )


def build_index(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Index Rules",
        [
            "List every attachment once, identify its purpose, and map it to the exact appeal statement "
            "it supports. Use final filenames that are readable outside the case-management system.",
        ],
    )
    add_grid(
        doc,
        ["No.", "Filename", "Source", "Purpose", "Claim IDs", "Reviewed"],
        [
            ["01", "01_Appeal.pdf", d.value("seller"), "Manual review request and action summary", "All", "Yes" if d.mode == "example" else "[YES/NO]"],
            ["02", "02_Claim_Register.xlsx", d.value("seller"), "Claim-to-evidence mapping", "All", "Yes" if d.mode == "example" else "[YES/NO]"],
            ["03", "03_Manufacturer_Declaration.pdf", d.value("manufacturer"), "Product and component mapping", "C-001", "Yes" if d.mode == "example" else "[YES/NO]"],
            ["04", "04_Test_Report.pdf", d.value("laboratory"), "Technical result", "C-002", "Yes" if d.mode == "example" else "[YES/NO]"],
            ["05", "05_Processing_Summary.xlsx", "Amazon", "SKU-level catalog processing result", "C-003", "Yes" if d.mode == "example" else "[YES/NO]"],
        ],
        [0.42, 1.7, 1.25, 2.0, 0.72, 0.72],
    )
    add_section(
        doc,
        "File Naming Standard",
        [
            "Use a two-digit order, concise document name, variation or product reference when needed, "
            "and version date. Do not use vague names such as image1, final-final, or document copy.",
        ],
    )


def build_checklist(doc: Document, d: CaseData) -> None:
    add_section(
        doc,
        "Evidence Gate",
        bullets=[
            "Every material statement has a named source and locator.",
            "Manufacturer and technical evidence map to every affected product.",
            "Limitations and unresolved questions are disclosed.",
            "All required signatures, stamps, dates, and authorities are present.",
        ],
    )
    add_section(
        doc,
        "Catalog And Content Gate",
        bullets=[
            "Affected ASINs, SKUs, variations, and marketplaces are consistent.",
            "Only intended SKUs and fields appear in correction files.",
            "Successful processing is supported by Amazon's processing summary.",
            "Frontend and backend corrections are verified after processing.",
            "Packaging, manual, media, and QR-linked content match the approved standard.",
        ],
    )
    add_section(
        doc,
        "Document Hygiene Gate",
        bullets=[
            "No unresolved placeholders, comments, tracked changes, hidden text, or personal metadata.",
            "DOCX and PDF versions render cleanly.",
            "The evidence index lists every attachment.",
            "The appeal does not overstate regulatory or technical conclusions.",
        ],
    )
    add_grid(
        doc,
        ["Review stage", "Reviewer", "Status", "Date", "Notes"],
        [
            ["Product accuracy", "[NAME]" if d.mode == "template" else "Example Product Owner", "Complete" if d.mode == "example" else "[STATUS]", "[DATE]" if d.mode == "template" else "February 2, 2026", ""],
            ["Regulated content", d.value("owner"), "Complete" if d.mode == "example" else "[STATUS]", "[DATE]" if d.mode == "template" else "February 3, 2026", ""],
            ["Victor signoff", "Victor Uhl", "Required", "[DATE]" if d.mode == "template" else "Not granted in fictional example", "Required before real submission"],
        ],
        [1.35, 1.65, 1.05, 1.45, 2.0],
    )
    add_callout(
        doc,
        "Submission Rule",
        "The pack may be labeled Approved for manual submission only after the validator passes and "
        "Victor's approval is recorded. Submission itself remains a separate action.",
    )


BUILDERS: dict[str, Callable[[Document, CaseData], None]] = {
    "01-appeal": build_appeal,
    "02-technical-evidence-and-claim-register": build_evidence,
    "03-manufacturer-declaration": build_declaration,
    "04-catalog-correction-and-processing-proof": build_catalog,
    "05-compliance-sop": build_sop,
    "06-training-guide": build_training,
    "07-manual-review": build_manual,
    "08-packaging-and-label-review": build_packaging,
    "09-evidence-index": build_index,
    "10-submission-checklist": build_checklist,
}


def build_document(module: str, title: str, data: CaseData, output: Path) -> None:
    doc = Document()
    configure_document(doc, data)
    add_title(doc, title, module, data)
    BUILDERS[module](doc, data)
    doc.add_section(WD_SECTION.CONTINUOUS)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_library(output_dir: Path) -> None:
    for mode, folder in (("template", "01-clean-templates"), ("example", "02-fictional-worked-example")):
        data = CaseData(mode)
        target = output_dir / folder
        for module, title in MODULES:
            build_document(module, title, data, target / f"{module}.docx")

    (output_dir / "REFERENCE-NOTE.txt").write_text(
        "This library is a reusable operating reference for serious Amazon product suppressions.\n"
        "The clean templates contain intentional placeholders. The worked example is entirely fictional.\n"
        "Neither set is legal advice or regulatory precedent. Real submissions require current evidence,\n"
        "current first-party Amazon requirements, and Victor's final troubleshooting signoff.\n",
        encoding="utf-8",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    build_library(args.output_dir)
    print(args.output_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
